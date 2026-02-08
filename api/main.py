"""
FastAPI Backend for Data Analysis Agent
Vercel Serverless Functions Compatible
"""

from fastapi import FastAPI, File, UploadFile, HTTPException, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, StreamingResponse
import pandas as pd
import numpy as np
import io
import json
import sys
from pathlib import Path
from typing import Optional, List, Dict, Any
import uuid
import tempfile
import os

# Add parent directory to path to import src module
parent_dir = str(Path(__file__).parent.parent)
if parent_dir not in sys.path:
    sys.path.insert(0, parent_dir)

from src.eda_agent import EDAAgent
from src.schema_compressor import SchemaCompressor

# File processors
try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    import docx
except ImportError:
    docx = None

try:
    import openpyxl
except ImportError:
    openpyxl = None

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

app = FastAPI(title="Data Analysis Agent API", version="2.0.0")

# CORS configuration
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Update with your frontend URL in production
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# In-memory storage for demo (use Redis/Database in production)
sessions: Dict[str, Dict[str, Any]] = {}


@app.get("/")
async def root():
    """Health check endpoint"""
    return {"status": "healthy", "service": "Data Analysis Agent API"}


@app.get("/api/health")
async def health():
    """API health check"""
    return {"status": "ok", "version": "2.0.0"}


def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text from PDF file"""
    if PyPDF2 is None:
        raise HTTPException(status_code=400, detail="PDF processing not available. Install PyPDF2.")
    
    pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text() + "\n"
    return text


def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from DOCX file"""
    if docx is None:
        raise HTTPException(status_code=400, detail="DOCX processing not available. Install python-docx.")
    
    doc = docx.Document(io.BytesIO(file_content))
    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    return text


def process_file_to_dataframe(file_content: bytes, filename: str) -> pd.DataFrame:
    """
    Process uploaded file and convert to DataFrame
    Supports: CSV, XLSX, XLS, JSON, PDF, DOCX, TXT
    """
    extension = filename.lower().split('.')[-1]
    
    try:
        if extension == 'csv':
            df = pd.read_csv(io.BytesIO(file_content))
        
        elif extension in ['xlsx', 'xls']:
            df = pd.read_excel(io.BytesIO(file_content))
        
        elif extension == 'json':
            df = pd.read_json(io.BytesIO(file_content))
        
        elif extension == 'pdf':
            # Use pdfplumber for better table extraction
            if pdfplumber is not None:
                try:
                    pdf_file = io.BytesIO(file_content)
                    
                    with pdfplumber.open(pdf_file) as pdf:
                        all_tables = []
                        
                        # Extract tables from all pages
                        for page in pdf.pages:
                            tables = page.extract_tables()
                            if tables:
                                for table in tables:
                                    if table and len(table) > 0:
                                        all_tables.extend(table)
                        
                        # If we found tables, convert to DataFrame
                        if all_tables and len(all_tables) > 1:
                            # First row is header
                            headers = all_tables[0]
                            data = all_tables[1:]
                            
                            print(f"DEBUG: Headers = {headers}")
                            print(f"DEBUG: First data row = {data[0] if data else 'No data'}")
                            
                            # Check if all data is in first cell (merged cell issue)
                            if data and data[0][0] and '\n' in str(data[0][0]) and all(not cell or str(cell) == 'None' for cell in data[0][1:]):
                                print("⚠️ Detected merged cells - extracting text and using CSV format parser")
                                text_content = str(data[0][0])
                                lines = [line.strip() for line in text_content.split('\n') if line.strip()]
                                
                                # Each line pattern: Title [#] [Type] [Genre] [Country...] [Year] [Rating] [Duration]
                                # Try to parse using regex for known patterns
                                import re
                                parsed_rows = []
                                
                                for line in lines:
                                    # Pattern: starts with "Title" followed by number, then type (TV Show/Movie)
                                    match = re.match(r'(Title\s+\d+)\s+(TV Show|Movie)\s+(\w+)\s+([\w\s]+?)\s+(\d{4})\s+([\w-]+)\s+(.*)', line)
                                    if match:
                                        title, type_, genre, country, year, rating, duration = match.groups()
                                        parsed_rows.append([
                                            title.strip(),
                                            type_.strip(),
                                            genre.strip(),
                                            country.strip(),
                                            year.strip(),
                                            rating.strip(),
                                            duration.strip()
                                        ])
                                
                                if parsed_rows:
                                    df = pd.DataFrame(parsed_rows, columns=headers)
                                    print(f"✅ Parsed {len(df)} rows from merged cells into {len(df.columns)} columns")
                                    print(f"First row: {df.iloc[0].to_dict()}")
                                    return df
                                else:
                                    print("❌ Failed to parse merged cells with regex")
                            
                            # Clean None values and empty strings
                            headers = [str(h).strip() if h and str(h).strip() else f"Column_{i}" for i, h in enumerate(headers)]
                            cleaned_data = []
                            for row in data:
                                # Make sure row has same number of elements as headers
                                if len(row) < len(headers):
                                    row = row + [''] * (len(headers) - len(row))
                                elif len(row) > len(headers):
                                    row = row[:len(headers)]
                                
                                cleaned_row = [str(cell).strip() if cell and str(cell).strip() not in ['None', 'nan', ''] else "" for cell in row]
                                cleaned_data.append(cleaned_row)
                            
                            df = pd.DataFrame(cleaned_data, columns=headers)
                            print(f"✅ pdfplumber extracted table with {len(df.columns)} columns, {len(df)} rows")
                            print(f"Column names: {list(df.columns)}")
                            return df
                except Exception as e:
                    print(f"pdfplumber extraction failed: {e}")
            
            # Fallback to text extraction
            text = extract_text_from_pdf(file_content)
            lines = [line.strip() for line in text.split('\n') if line.strip()]
            
            if not lines:
                raise HTTPException(status_code=400, detail="PDF appears to be empty")
            
            # Try fixed-width format
            try:
                df = pd.read_fwf(io.StringIO('\n'.join(lines)), infer_nrows=min(200, len(lines)))
                
                if len(df.columns) > 1 and len(df) > 0:
                    df.columns = [str(col).strip() for col in df.columns]
                    print(f"✅ Fixed-width parsing: {len(df.columns)} columns")
                    return df
            except Exception as e:
                print(f"Fixed-width failed: {e}")
            
            # Last resort: single column
            print(f"⚠️ PDF fallback to single column")
            df = pd.DataFrame({'content': lines})
        
        elif extension in ['docx', 'doc']:
            text = extract_text_from_docx(file_content)
            lines = [line.strip() for line in text.split('\n') if line.strip()]
            
            if not lines:
                raise HTTPException(status_code=400, detail="Document appears to be empty")
            
            # Try fixed-width format for tables
            try:
                df = pd.read_fwf(io.StringIO('\n'.join(lines)), infer_nrows=200)
                if len(df.columns) > 1 and len(df) > 0:
                    df.columns = [str(col).strip() for col in df.columns]
                    return df
            except:
                pass
            
            # Fallback to single column
            df = pd.DataFrame({'content': lines})
        
        elif extension == 'txt':
            text = file_content.decode('utf-8')
            lines = [line.strip() for line in text.split('\n') if line.strip()]
            
            if not lines:
                raise HTTPException(status_code=400, detail="Text file appears to be empty")
            
            first_line = lines[0]
            
            # Try CSV format first
            if ',' in first_line:
                try:
                    df = pd.read_csv(io.StringIO(text))
                    if len(df.columns) > 1:
                        return df
                except:
                    pass
            
            # Try tab-separated
            if '\t' in first_line:
                try:
                    df = pd.read_csv(io.StringIO(text), sep='\t')
                    if len(df.columns) > 1:
                        return df
                except:
                    pass
            
            # Try fixed-width format
            try:
                df = pd.read_fwf(io.StringIO(text), infer_nrows=200)
                if len(df.columns) > 1 and len(df) > 0:
                    df.columns = [str(col).strip() for col in df.columns]
                    return df
            except:
                pass
            
            # Fallback to single column
            df = pd.DataFrame({'content': lines})
        
        else:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported file format: {extension}"
            )
        
        return df
    
    except Exception as e:
        print(f"❌ Upload error: {str(e)}")
        import traceback
        traceback.print_exc()
        raise HTTPException(
            status_code=400,
            detail=f"Error processing file: {str(e)}"
        )


@app.post("/api/upload")
async def upload_file(file: UploadFile = File(...)):
    """
    Upload and process data file
    Returns session_id and initial data preview
    """
    try:
        print(f"📤 Upload request - Filename: {file.filename}, Content-Type: {file.content_type}")
        
        # Read file content
        content = await file.read()
        print(f"📦 File size: {len(content)} bytes")
        
        # Process file to DataFrame
        df = process_file_to_dataframe(content, file.filename)
        
        # Create session
        session_id = str(uuid.uuid4())
        
        # Initialize EDA Agent
        agent = EDAAgent(df, name="Web Agent")
        
        # Store session data
        sessions[session_id] = {
            'dataframe': df,
            'agent': agent,
            'filename': file.filename,
            'original_shape': df.shape,
            'current_sort': None,
            'current_filters': None
        }
        
        # Prepare response
        preview_data = df.head(100).to_dict('records')
        columns = [{'field': col, 'headerName': col, 'type': str(df[col].dtype)} 
                   for col in df.columns]
        
        return JSONResponse({
            'session_id': session_id,
            'filename': file.filename,
            'shape': {'rows': df.shape[0], 'columns': df.shape[1]},
            'columns': columns,
            'data': preview_data,
            'dtypes': df.dtypes.astype(str).to_dict(),
            'summary': agent.get_schema_context()[:500]
        })
    
    except HTTPException as he:
        raise he
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/data/{session_id}")
async def get_data(
    session_id: str,
    page: int = 0,
    page_size: int = 100,
    sort_column: Optional[str] = None,
    sort_order: Optional[str] = 'asc'
):
    """
    Get paginated and sorted data
    """
    if session_id not in sessions:
        raise HTTPException(status_code=404, detail="Session not found")
    
    df = sessions[session_id]['dataframe'].copy()
    
    # Apply sorting
    if sort_column and sort_column in df.columns:
        ascending = sort_order == 'asc'
        df = df.sort_values(by=sort_column, ascending=ascending)
    
    # Pagination
    start_idx = page * page_size
    end_idx = start_idx + page_size
    paginated_df = df.iloc[start_idx:end_idx]
    
    return JSONResponse({
        'data': paginated_df.to_dict('records'),
        'total_rows': len(df),
        'page': page,
        'page_size': page_size
    })


@app.post("/api/filter/{session_id}")
async def filter_data(
    session_id: str,
    filters: Dict[str, Any]
):
    """
    Apply filters to data
    filters format: {
        'column_name': {'operator': 'eq|ne|gt|lt|contains', 'value': '...'},
        ...
    }
    """
    if session_id not in sessions:
        raise HTTPException(status_code=404, detail="Session not found")
    
    df = sessions[session_id]['dataframe'].copy()
    
    # Apply filters
    for column, filter_config in filters.items():
        if column not in df.columns:
            continue
        
        operator = filter_config.get('operator', 'eq')
        value = filter_config.get('value')
        
        # Convert value to appropriate type
        col_dtype = df[column].dtype
        try:
            if pd.api.types.is_numeric_dtype(col_dtype):
                value = pd.to_numeric(value)
            elif pd.api.types.is_bool_dtype(col_dtype):
                value = str(value).lower() in ['true', '1', 'yes']
        except (ValueError, TypeError):
            pass  # Keep as string
        
        if operator == 'eq':
            df = df[df[column] == value]
        elif operator == 'ne':
            df = df[df[column] != value]
        elif operator == 'gt':
            df = df[df[column] > value]
        elif operator == 'lt':
            df = df[df[column] < value]
        elif operator == 'gte':
            df = df[df[column] >= value]
        elif operator == 'lte':
            df = df[df[column] <= value]
        elif operator == 'contains':
            df = df[df[column].astype(str).str.contains(str(value), case=False, na=False)]
    
    # Store current filters
    sessions[session_id]['current_filters'] = filters
    
    return JSONResponse({
        'data': df.to_dict('records'),
        'filtered_rows': len(df),
        'total_rows': sessions[session_id]['dataframe'].shape[0]
    })


@app.post("/api/export/{session_id}")
async def export_data(
    session_id: str,
    filters: Optional[Dict[str, Any]] = None,
    columns: Optional[List[str]] = None,
    format: str = 'csv',
    sort_column: Optional[str] = None,
    sort_order: str = 'asc'
):
    """
    Export filtered and sorted data in specified format
    """
    if session_id not in sessions:
        raise HTTPException(status_code=404, detail="Session not found")
    
    df = sessions[session_id]['dataframe'].copy()
    
    # Apply filters if provided (use current filters if not specified)
    filters_to_apply = filters or sessions[session_id].get('current_filters')
    if filters_to_apply:
        for column, filter_config in filters_to_apply.items():
            if column not in df.columns:
                continue
            
            operator = filter_config.get('operator', 'eq')
            value = filter_config.get('value')
            
            # Convert value to appropriate type
            col_dtype = df[column].dtype
            try:
                if pd.api.types.is_numeric_dtype(col_dtype):
                    value = pd.to_numeric(value)
                elif pd.api.types.is_bool_dtype(col_dtype):
                    value = str(value).lower() in ['true', '1', 'yes']
            except (ValueError, TypeError):
                pass
            
            if operator == 'eq':
                df = df[df[column] == value]
            elif operator == 'ne':
                df = df[df[column] != value]
            elif operator == 'gt':
                df = df[df[column] > value]
            elif operator == 'lt':
                df = df[df[column] < value]
            elif operator == 'gte':
                df = df[df[column] >= value]
            elif operator == 'lte':
                df = df[df[column] <= value]
            elif operator == 'contains':
                df = df[df[column].astype(str).str.contains(str(value), case=False, na=False)]
    
    # Apply sorting (use current sort if not specified)
    if sort_column and sort_column in df.columns:
        df = df.sort_values(by=sort_column, ascending=sort_order.lower() == 'asc')
    elif sessions[session_id].get('current_sort'):
        sort_info = sessions[session_id]['current_sort']
        if sort_info['column'] in df.columns:
            df = df.sort_values(by=sort_info['column'], ascending=sort_info['order'].lower() == 'asc')
    
    # Select columns if specified
    if columns:
        df = df[[col for col in columns if col in df.columns]]
    
    # Export in requested format
    if format == 'csv':
        output = io.StringIO()
        df.to_csv(output, index=False)
        output.seek(0)
        return StreamingResponse(
            io.BytesIO(output.getvalue().encode()),
            media_type="text/csv",
            headers={"Content-Disposition": f"attachment; filename=export.csv"}
        )
    
    elif format == 'xlsx':
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, index=False)
        output.seek(0)
        return StreamingResponse(
            output,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": f"attachment; filename=export.xlsx"}
        )
    
    elif format == 'json':
        json_str = df.to_json(orient='records')
        return StreamingResponse(
            io.BytesIO(json_str.encode()),
            media_type="application/json",
            headers={"Content-Disposition": f"attachment; filename=export.json"}
        )
    
    elif format == 'pdf':
        # Create simple PDF with table data
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import letter, A4, landscape
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
        from reportlab.lib.units import inch
        
        output = io.BytesIO()
        
        # Use landscape for wide tables
        pagesize = landscape(A4) if len(df.columns) > 6 else A4
        doc = SimpleDocTemplate(output, pagesize=pagesize, 
                                leftMargin=0.5*inch, rightMargin=0.5*inch,
                                topMargin=0.5*inch, bottomMargin=0.5*inch)
        elements = []
        
        # Prepare data - ensure proper string conversion
        header = [str(col) for col in df.columns.tolist()]
        rows = []
        
        # Convert each row properly to avoid misalignment
        for idx, row in df.head(1000).iterrows():
            row_data = [str(val) if pd.notna(val) else '' for val in row.values]
            rows.append(row_data)
        
        # Combine header and data
        table_data = [header] + rows
        
        # Calculate column widths dynamically
        num_cols = len(df.columns)
        available_width = (pagesize[0] - inch) / num_cols
        col_widths = [available_width] * num_cols
        
        # Create table with column widths
        table = Table(table_data, colWidths=col_widths, repeatRows=1)
        # Style the table
        table.setStyle(TableStyle([
            # Header styling
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#667eea')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 7),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
            ('TOPPADDING', (0, 0), (-1, 0), 8),
            
            # Data rows styling
            ('BACKGROUND', (0, 1), (-1, -1), colors.white),
            ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
            ('ALIGN', (0, 1), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 1), (-1, -1), 6),
            ('TOPPADDING', (0, 1), (-1, -1), 4),
            ('BOTTOMPADDING', (0, 1), (-1, -1), 4),
            ('LEFTPADDING', (0, 0), (-1, -1), 4),
            ('RIGHTPADDING', (0, 0), (-1, -1), 4),
            
            # Grid
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            
            # Alternating row colors
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f5f5f5')])
        ]))
        
        elements.append(table)
        doc.build(elements)
        output.seek(0)
        
        return StreamingResponse(
            output,
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename=export.pdf"}
        )
    
    elif format == 'docx':
        # Create DOCX file with table
        if docx is None:
            raise HTTPException(status_code=400, detail="python-docx not installed")
        
        doc = docx.Document()
        doc.add_heading('Data Export', 0)
        
        # Add table
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = 'Light Grid Accent 1'
        
        # Header row
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(df.columns):
            hdr_cells[i].text = str(col)
        
        # Data rows (limit to 1000 for performance)
        for _, row in df.head(1000).iterrows():
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                row_cells[i].text = str(value)
        
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        
        return StreamingResponse(
            output,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=export.docx"}
        )
    
    else:
        raise HTTPException(status_code=400, detail=f"Unsupported format: {format}")


@app.post("/api/analyze/{session_id}")
async def analyze_data(session_id: str, query: str = Form(...)):
    """
    Perform AI-powered analysis on the data
    """
    if session_id not in sessions:
        raise HTTPException(status_code=404, detail="Session not found")
    
    agent = sessions[session_id]['agent']
    
    # Get basic insights
    insights = agent.get_basic_insights()
    
    return JSONResponse({
        'insights': insights,
        'schema': agent.get_schema_context()
    })


@app.get("/api/statistics/{session_id}")
async def get_statistics(session_id: str):
    """Get statistical summary of the data"""
    if session_id not in sessions:
        raise HTTPException(status_code=404, detail="Session not found")
    
    df = sessions[session_id]['dataframe']
    
    # Numeric statistics
    numeric_stats = df.describe().to_dict()
    
    # Categorical statistics
    categorical_stats = {}
    for col in df.select_dtypes(include=['object']).columns:
        categorical_stats[col] = {
            'unique_values': int(df[col].nunique()),
            'top_value': str(df[col].mode()[0]) if len(df[col].mode()) > 0 else None,
            'top_frequency': int(df[col].value_counts().iloc[0]) if len(df[col]) > 0 else 0
        }
    
    # Missing values
    missing_values = df.isnull().sum().to_dict()
    
    return JSONResponse({
        'numeric_statistics': numeric_stats,
        'categorical_statistics': categorical_stats,
        'missing_values': missing_values,
        'total_rows': len(df),
        'total_columns': len(df.columns)
    })


# Vercel serverless function handler
def handler(request):
    """Vercel serverless function handler"""
    return app(request)
